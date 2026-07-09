#!/usr/bin/env python3
"""Generate LEO Tickets interview guide as Word document."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

OUTPUT = Path(__file__).resolve().parent.parent / "LEO_Tickets_Interview_Guide.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LEO Tickets — Interview Deep Dive", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading(doc, "Elevator Pitch (30 seconds)", 1)
    add_para(
        doc,
        "LEO Tickets is a full-stack event ticketing web app for Yale/LEO social events. "
        "Guests log in with Google OAuth, get a QR ticket if they are on a guest list, and "
        "door staff use a mobile scanner that looks up Yale directory data and logs entries. "
        "Admins manage the guest list, blacklist, event title, and scan logs via a web panel. "
        "I rebuilt it from a flat-file prototype (text files + Excel) into a production-style "
        "app: PostgreSQL, RBAC, REST API, Docker, CI, deployed on Render + Neon.",
    )

    add_heading(doc, "Problem & Constraints", 1)
    add_table(
        doc,
        ["Need", "Solution"],
        [
            ["Verify identity", "Google OAuth (Yale emails)"],
            ["Control who gets tickets", "Allowlist + silent blacklist"],
            ["Fast door check", "QR scan → decode email → lookup name/photo"],
            ["Event-night ops", "Admin panel, no redeploy for guest list changes"],
            ["Small team, low budget", "Free tier: Render + Neon, no AWS"],
        ],
    )

    add_heading(doc, "Architecture", 1)
    add_para(doc, "Pattern: Flask application factory + blueprints + thin routes, fat services.")
    add_bullets(
        doc,
        [
            "app/__init__.py — create_app(), register blueprints",
            "app/config.py — env-based config + production validation",
            "app/models.py — SQLAlchemy ORM",
            "app/auth/ — OAuth, RBAC, decorators",
            "app/services/ — business logic (events, scans, users, qr)",
            "app/routes/ — server-rendered HTML (main, admin)",
            "app/api/v1/ — REST JSON API",
        ],
    )

    add_heading(doc, "Technology Stack", 1)
    add_heading(doc, "Backend", 2)
    add_table(
        doc,
        ["Tech", "Role"],
        [
            ["Python 3.11", "Runtime"],
            ["Flask 3", "Web framework, sessions, Jinja templates"],
            ["SQLAlchemy 2 + Flask-SQLAlchemy", "ORM"],
            ["Flask-Migrate (Alembic)", "Schema migrations"],
            ["psycopg2", "PostgreSQL driver"],
            ["Gunicorn", "Production WSGI server"],
            ["Authlib", "OAuth 2.0 client for Google"],
            ["google-api-python-client", "Fetch user profile after OAuth"],
            ["qrcode + Pillow", "Generate QR PNG as base64"],
        ],
    )
    add_heading(doc, "Data & Infrastructure", 2)
    add_table(
        doc,
        ["Tech", "Role"],
        [
            ["PostgreSQL 16", "Primary datastore"],
            ["Neon", "Managed Postgres (serverless, free tier)"],
            ["Docker + docker-compose", "Local Postgres + optional full stack"],
            ["Render", "Container hosting from GitHub"],
            ["GitHub Actions", "CI: migrations + smoke import"],
        ],
    )
    add_heading(doc, "Frontend", 2)
    add_bullets(
        doc,
        [
            "Jinja2 templates — SSR admin, ticket, scanner pages",
            "html5-qrcode — browser camera QR scanning",
            "Vanilla JS — scanner fetch, admin image toggle",
            "No React/SPA — intentional for a small ops tool",
        ],
    )

    add_heading(doc, "Core User Flows", 1)
    add_heading(doc, "1. Guest Ticket (GET /)", 2)
    add_bullets(
        doc,
        [
            "@login_required → redirect to Google OAuth if no session",
            "Load email from Google userinfo API",
            "is_attendee(email) → must be on allowlist AND not blacklisted",
            "Encode email with custom cipher → generate QR → render ticket page",
        ],
    )
    add_para(
        doc,
        "Interview point: Superadmin can access /admin but still needs allowlist for /. "
        "Authorization is route-specific, not one global logged-in = full access.",
        bold=False,
    )

    add_heading(doc, "2. Scanner (GET /scanner, POST /scanner_result)", 2)
    add_bullets(
        doc,
        [
            "Scanner page is public (kiosk-friendly)",
            "POST requires scans:write via X-API-Key or Google session with scanner/admin role",
            "process_scan(): decode QR → email, dedup within 60s, insert ScanLog, lookup YaleStudent",
            "Respects EventSettings.image_visible toggle",
        ],
    )

    add_heading(doc, "3. Admin (GET /admin)", 2)
    add_bullets(
        doc,
        [
            "Must be logged in + have event:write permission",
            "CRUD on guest list, blacklist, ticket title via form POSTs",
            "Blacklist silently removes from allowlist; blocked users see same error as not on list",
        ],
    )

    add_heading(doc, "Auth & RBAC Design", 1)
    add_para(doc, "Two auth mechanisms, one permission system:")
    add_table(
        doc,
        ["Method", "Use case"],
        [
            ["Google OAuth session", "Guests, admins (browser cookies)"],
            ["API keys (X-API-Key / Bearer)", "Scanner kiosks, automation"],
        ],
    )
    add_bullets(
        doc,
        [
            "Roles: scanner, admin, superadmin",
            "Permissions: string-based (event:write, scans:read, etc.)",
            "Mapping: static dict in permissions.py",
            "Superadmin bootstrap: SUPERADMIN_EMAIL env var",
            "Staff admins: seeded into users + role_assignments",
            "AuthContext on Flask g — unified has_permission() for web + API",
            "OAuth: CSRF state validated; tokens in signed session; no_cache on auth routes",
        ],
    )

    add_heading(doc, "Database Schema", 1)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["event_settings", "Single-row event config (title, image toggle)"],
            ["allowed_emails", "Guest allowlist"],
            ["blacklisted_emails", "Silent deny list"],
            ["users + role_assignments", "RBAC staff"],
            ["yale_students", "~47k rows — directory lookup for scanner"],
            ["scan_logs", "Timestamped entry log"],
        ],
    )
    add_bullets(
        doc,
        [
            "Normalized allowlist — indexable, API-friendly",
            "Separate blacklist table — independent policy",
            "YaleStudent PK = email — O(1) lookup on scan",
            "UTC timestamps on all datetime columns",
            "3 Alembic migrations: initial schema, RBAC users, blacklist",
        ],
    )

    add_heading(doc, "REST API (/api/v1)", 1)
    add_table(
        doc,
        ["Endpoint", "Permission"],
        [
            ["GET /health", "Public"],
            ["GET /me", "Any auth"],
            ["GET/PATCH /event", "event:read / event:write"],
            ["GET/PUT /allowed-emails", "allowlist:read / write"],
            ["GET/PUT /blacklisted-emails", "allowlist:read / write"],
            ["GET /ticket", "Session + allowlist"],
            ["POST /scans, GET /scans", "scans:write / read"],
            ["GET/POST/DELETE /users", "superadmin only"],
        ],
    )
    add_para(doc, 'Error shape: { "error": { "code": "FORBIDDEN", "message": "..." } }')

    add_heading(doc, "QR Encoding", 1)
    add_bullets(
        doc,
        [
            "Custom substitution cipher on email — obfuscation, NOT encryption",
            "Hides raw email from casual QR readers; backward compatible with legacy app",
            "Not cryptographically secure — anyone with source can decode",
            "Better alternative: HMAC-signed JWT or Fernet token with server secret + expiry",
        ],
    )

    add_heading(doc, "Migration Story (Legacy → Modern)", 1)
    add_para(doc, "Before: Monolithic app.py, flat files, hardcoded secrets, Render deploy.")
    add_para(doc, "After:")
    add_bullets(
        doc,
        [
            "Application factory + blueprints",
            "Postgres as source of truth",
            "Env-based secrets + production validation",
            "Service layer extracted from routes",
            "Seed script imports legacy files → DB",
            "Sensitive/PII data gitignored; example files in repo",
            "Git history rewritten before going public",
        ],
    )

    add_heading(doc, "Deployment & Ops", 1)
    add_para(doc, "GitHub (main) → Render (Docker) → Neon (Postgres)")
    add_bullets(
        doc,
        [
            "Dockerfile: flask db upgrade && gunicorn on startup",
            "PORT env for Render compatibility",
            "Seeding: manual from laptop (scripts/seed_database.py)",
            "CI: GitHub Actions runs migrations + create_app() smoke test",
            "Render auto-deploys on push; no separate CD pipeline",
        ],
    )

    add_heading(doc, "Tradeoffs", 1)
    add_table(
        doc,
        ["Decision", "Pro", "Con"],
        [
            ["SSR + Jinja vs SPA", "Simple, fast to ship", "Less interactive UX"],
            ["Static RBAC map", "Simple, auditable", "Deploy to change permissions"],
            ["API key in scanner HTML", "Kiosk without login", "Key visible in page source"],
            ["Full replace for allowlist", "Simple admin UX", "Not ideal at huge scale"],
            ["Single event_settings row", "One event at a time", "No multi-event support"],
            ["47k Yale rows in Postgres", "Fast lookup", "Large seed time, storage"],
            ["Smoke test only in CI", "Shipped quickly", "Regression risk"],
        ],
    )

    add_heading(doc, "Likely Interview Questions & Answers", 1)

    qa = [
        (
            "Why PostgreSQL over flat files?",
            "Concurrent writes (multiple scanners), relational integrity, indexed lookups, "
            "migrations, and a path to proper auth/RBAC. Excel/text files don't scale for "
            "event-night concurrent scans.",
        ),
        (
            "How does auth work?",
            "Dual path: Google OAuth for humans (session cookies signed with SECRET_KEY), "
            "API keys for machines. Both resolve to AuthContext with roles; permissions "
            "checked via decorators before handlers run.",
        ),
        (
            "How would you scale for 5,000 guests at the door?",
            "Connection pooling (Neon pooler), rate limiting on /scanner_result, Redis for "
            "dedup, paid Render tier to avoid cold starts, horizontal scaling with stateless "
            "app servers.",
        ),
        (
            "Security concerns?",
            "QR obfuscation isn't encryption; scanner API key in client; OAuth state for CSRF; "
            "secrets in env vars; silent blacklist for privacy. Would add signed QR tokens with "
            "expiry, HTTPS-only cookies, audit logs.",
        ),
        (
            "Why a service layer?",
            "Routes stay thin (HTTP only); business logic reusable by web + API; easier to "
            "test; mirrors larger backend structure.",
        ),
        (
            "What was the hardest part?",
            "Migrating legacy behavior without breaking event workflow; seeding 47k directory "
            "rows to remote Neon; OAuth redirect URI mismatches; separating PII from public repo.",
        ),
    ]
    for q, a in qa:
        add_para(doc, f"Q: {q}", bold=True)
        add_para(doc, f"A: {a}")
        doc.add_paragraph()

    add_heading(doc, "What to Improve / Add", 1)
    add_heading(doc, "High Impact", 2)
    add_bullets(
        doc,
        [
            "Automated tests — pytest for is_attendee, RBAC, process_scan dedup, API auth",
            "Signed QR tokens — HMAC(email + event_id + expiry) instead of substitution cipher",
            "Rate limiting — Flask-Limiter on scan endpoint",
            "Seed in CI/CD or admin UI — avoid manual laptop seed to production",
            "Scanner auth redesign — short-lived session instead of API key in HTML",
        ],
    )
    add_heading(doc, "Product", 2)
    add_bullets(
        doc,
        [
            "Multi-event support — events table with FKs",
            "Admin UI for RBAC — manage scanner operators without DB seed",
            "Real-time scan dashboard — WebSockets or SSE",
            "Export scan log — CSV download from admin",
            "Email notifications — ticket confirmation, waitlist",
        ],
    )
    add_heading(doc, "Engineering / Ops", 2)
    add_bullets(
        doc,
        [
            "Structured logging + monitoring (Sentry, health checks with DB ping)",
            "Pagination on admin scan log view",
            "Staging environment — separate Neon branch + Render preview",
            "Pre-commit hooks — secret scanning, ruff/black",
            "OpenAPI/Swagger for /api/v1",
            "Redis caching for YaleStudent hot path",
            "GDPR/privacy — retention policy for scan logs and directory data",
        ],
    )
    add_heading(doc, "Security Hardening", 2)
    add_bullets(
        doc,
        [
            "HttpOnly, Secure, SameSite cookie flags in production",
            "CORS policy if SPA clients added",
            "Audit log for admin actions",
            "Remove hardcoded cipher — env-derived key for QR signing",
        ],
    )

    add_heading(doc, "Resume One-Liner", 1)
    add_para(
        doc,
        "LEO Tickets — Event ticketing platform with OAuth guest verification, QR-based entry "
        "scanning, and admin allowlist management. Rebuilt legacy flat-file app into "
        "PostgreSQL-backed Flask service with RBAC, REST API, Docker, and CI/CD; deployed "
        "to production on Render.",
    )

    add_heading(doc, "Technologies One-Liner", 1)
    add_para(
        doc,
        "Python/Flask backend with SQLAlchemy and PostgreSQL, Google OAuth 2.0, role-based "
        "access control, a REST API, Dockerized deployment on Render with Neon Postgres, "
        "GitHub Actions CI, and client-side QR scanning via the browser MediaDevices API.",
    )

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
